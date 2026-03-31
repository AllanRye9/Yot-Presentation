#!/usr/bin/env python3
"""
Tests for the Yot-Presentation web application (web/app.py).

Run with:  python -m pytest tests/test_web_app.py -v
"""

import io
import json
import sys
from pathlib import Path

import pytest

# Allow importing from web/
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from web.app import (
    app,
    match_command,
    convert_text,
    convert_image,
    convert_csv,
    _extract_keywords,
    _extractive_summary,
    _simple_sentiment,
    _estimate_reading_time,
    _analyze_chart_image,
)


# ─── fixtures ────────────────────────────────────────────────────────────


@pytest.fixture
def client():
    app.config["TESTING"] = True
    with app.test_client() as c:
        yield c


# ─── command matching (mirrors original v5.3.1 patterns) ─────────────────


class TestMatchCommand:
    """Validate every voice-command action from the original application."""

    def _check(self, text, expected_action, **extra):
        result = match_command(text)
        assert result["action"] == expected_action, (
            f'"{text}" → {result["action"]!r}, expected {expected_action!r}'
        )
        for k, v in extra.items():
            assert result[k] == v, f"Expected {k}={v!r}, got {result[k]!r}"

    # navigation
    def test_next_slide_english(self):
        for phrase in ["next", "forward", "advance", "go forward", "next slide"]:
            self._check(phrase, "next_slide")

    def test_prev_slide_english(self):
        for phrase in ["previous", "back", "go back", "return"]:
            self._check(phrase, "prev_slide")

    def test_jump_slide_extracts_number(self):
        self._check("slide 5", "jump_slide", slide=5)
        self._check("go to 10", "jump_slide", slide=10)
        self._check("jump to 3", "jump_slide", slide=3)
        self._check("page 7", "jump_slide", slide=7)

    def test_first_last_slide(self):
        self._check("first slide", "first_slide")
        self._check("last slide", "last_slide")
        self._check("final slide", "last_slide")

    # presentation control
    def test_start_show(self):
        for phrase in ["start", "begin", "present", "start presentation"]:
            self._check(phrase, "start_show")

    def test_end_show(self):
        for phrase in ["end", "stop", "exit", "quit", "end show"]:
            self._check(phrase, "end_show")

    def test_blackout(self):
        for phrase in ["black screen", "blackout", "darken"]:
            self._check(phrase, "blackout")

    # zoom
    def test_zoom_in(self):
        for phrase in ["zoom in", "magnify", "enlarge"]:
            self._check(phrase, "zoom_in")

    def test_zoom_out(self):
        for phrase in ["zoom out", "shrink", "reduce"]:
            self._check(phrase, "zoom_out")

    def test_zoom_reset(self):
        for phrase in ["reset zoom", "normal size", "actual size"]:
            self._check(phrase, "zoom_reset")

    # annotations
    def test_pen_tool(self):
        for phrase in ["pen tool", "draw", "annotation"]:
            self._check(phrase, "pen_tool")

    def test_eraser(self):
        for phrase in ["eraser", "erase", "clear drawing"]:
            self._check(phrase, "eraser")

    def test_pointer(self):
        for phrase in ["pointer", "arrow", "cursor"]:
            self._check(phrase, "pointer")

    # fullscreen
    def test_fullscreen(self):
        self._check("fullscreen", "fullscreen")
        self._check("full screen", "fullscreen")

    # multi-language
    def test_spanish(self):
        self._check("siguiente", "next_slide")
        self._check("anterior", "prev_slide")
        self._check("diapositiva 3", "jump_slide", slide=3)

    def test_french(self):
        self._check("suivant", "next_slide")
        self._check("précédent", "prev_slide")

    def test_german(self):
        self._check("nächst", "next_slide")
        self._check("zurück", "prev_slide")

    def test_italian(self):
        self._check("prossimo", "next_slide")
        self._check("indietro", "prev_slide")

    def test_portuguese(self):
        self._check("próximo", "next_slide")
        self._check("voltar", "prev_slide")

    def test_chinese(self):
        self._check("下一张", "next_slide")
        self._check("上一张", "prev_slide")

    def test_japanese(self):
        self._check("次へ", "next_slide")
        self._check("戻る", "prev_slide")

    def test_unknown_returns_correct_structure(self):
        result = match_command("hello world")
        assert result["action"] == "unknown"
        assert result["confidence"] == 0


# ─── file conversion ─────────────────────────────────────────────────────


class TestConvertText:
    def test_splits_on_blank_lines(self):
        content = b"Slide One\n\nSlide Two\nBullet\n\nSlide Three"
        slides = convert_text(content)
        assert len(slides) == 3
        assert slides[0]["title"] == "Slide One"
        assert slides[1]["title"] == "Slide Two"
        assert slides[1]["bullets"][0]["text"] == "Bullet"

    def test_empty_content_gives_one_slide(self):
        slides = convert_text(b"   ")
        assert len(slides) == 1

    def test_slide_ids_are_sequential(self):
        slides = convert_text(b"A\n\nB\n\nC")
        assert [s["id"] for s in slides] == [1, 2, 3]

    def test_type_is_text(self):
        slides = convert_text(b"Hello world")
        assert slides[0]["type"] == "text"


class TestConvertImage:
    def test_returns_single_image_slide(self):
        fake_png = (
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
            + b"\x00" * 100  # not a real PNG, just testing structure
        )
        # Use a real JPEG from assets if available
        assets = Path(__file__).resolve().parents[1] / "assets"
        jpeg_path = assets / "a.jpeg"
        if jpeg_path.exists():
            img_bytes = jpeg_path.read_bytes()
        else:
            img_bytes = fake_png

        slides = convert_image(img_bytes, "test.jpg")
        assert len(slides) == 1
        assert slides[0]["type"] == "image"
        assert slides[0]["id"] == 1
        assert slides[0]["image"].startswith("data:image/jpeg;base64,")

    def test_png_mime_type(self):
        slides = convert_image(b"fake", "screenshot.png")
        assert slides[0]["image"].startswith("data:image/png;base64,")

    def test_title_is_filename_stem(self):
        slides = convert_image(b"fake", "my_photo.webp")
        assert slides[0]["title"] == "my_photo"


class TestConvertDocx:
    def test_heading_based_splitting(self):
        """Requires python-docx; creates a real DOCX in memory."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        from web.app import convert_docx

        doc = Document()
        doc.add_heading("Intro", 1)
        doc.add_paragraph("Introduction body.")
        doc.add_heading("Chapter 1", 1)
        doc.add_heading("Sub A", 2)
        doc.add_paragraph("Sub body.")
        buf = io.BytesIO()
        doc.save(buf)

        slides = convert_docx(buf.getvalue())
        assert len(slides) == 2
        assert slides[0]["title"] == "Intro"
        assert slides[1]["title"] == "Chapter 1"

    def test_fallback_chunking_when_no_headings(self):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        from web.app import convert_docx

        doc = Document()
        for i in range(20):
            doc.add_paragraph(f"Paragraph {i + 1}")
        buf = io.BytesIO()
        doc.save(buf)

        slides = convert_docx(buf.getvalue())
        assert len(slides) >= 2  # should be chunked


class TestConvertXlsx:
    def test_sheet_becomes_slide(self):
        try:
            import openpyxl
        except ImportError:
            pytest.skip("openpyxl not installed")

        from web.app import convert_xlsx

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sales"
        ws.append(["Product", "Revenue"])
        ws.append(["A", 1000])
        ws.append(["B", 2000])
        buf = io.BytesIO()
        wb.save(buf)

        slides = convert_xlsx(buf.getvalue())
        assert len(slides) == 1
        assert slides[0]["title"] == "Sales"
        assert slides[0]["headers"] == ["Product", "Revenue"]
        assert len(slides[0]["rows"]) == 2
        assert slides[0]["type"] == "table"


# ─── Flask API routes ─────────────────────────────────────────────────────


class TestFlaskRoutes:
    def test_index_returns_html(self, client):
        resp = client.get("/")
        assert resp.status_code == 200
        assert b"Yot-Presentation" in resp.data
        assert b"upload-screen" in resp.data

    def test_upload_no_file_400(self, client):
        resp = client.post("/upload")
        assert resp.status_code == 400
        data = json.loads(resp.data)
        assert "error" in data

    def test_upload_unsupported_type_400(self, client):
        resp = client.post(
            "/upload",
            data={"file": (io.BytesIO(b"data"), "file.xyz")},
            content_type="multipart/form-data",
        )
        assert resp.status_code == 400

    def test_upload_text_file(self, client):
        content = b"Section One\n\nSection Two\nDetails here"
        resp = client.post(
            "/upload",
            data={"file": (io.BytesIO(content), "doc.txt")},
            content_type="multipart/form-data",
        )
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert data["success"] is True
        assert data["total_slides"] >= 2
        assert data["filename"] == "doc.txt"

    def test_upload_jpeg(self, client):
        assets = Path(__file__).resolve().parents[1] / "assets" / "a.jpeg"
        if not assets.exists():
            pytest.skip("test asset not found")
        img_bytes = assets.read_bytes()
        resp = client.post(
            "/upload",
            data={"file": (io.BytesIO(img_bytes), "a.jpeg")},
            content_type="multipart/form-data",
        )
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert data["total_slides"] == 1
        assert data["slides"][0]["type"] == "image"

    def test_api_command_next(self, client):
        resp = client.post(
            "/api/command",
            data=json.dumps({"text": "next slide"}),
            content_type="application/json",
        )
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert data["action"] == "next_slide"

    def test_api_command_jump(self, client):
        resp = client.post(
            "/api/command",
            data=json.dumps({"text": "go to slide 7"}),
            content_type="application/json",
        )
        data = json.loads(resp.data)
        assert data["action"] == "jump_slide"
        assert data["slide"] == 7

    def test_api_command_unknown(self, client):
        resp = client.post(
            "/api/command",
            data=json.dumps({"text": "pizza please"}),
            content_type="application/json",
        )
        data = json.loads(resp.data)
        assert data["action"] == "unknown"

    def test_api_command_empty_body(self, client):
        resp = client.post(
            "/api/command",
            data=json.dumps({}),
            content_type="application/json",
        )
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert data["action"] == "unknown"


# ─── file management API ─────────────────────────────────────────────────


class TestFileManagement:
    """Tests for /api/files endpoints (file registry)."""

    def _upload_txt(self, client, content=b"Hello\n\nWorld", name="test.txt"):
        return client.post(
            "/upload",
            data={"file": (io.BytesIO(content), name)},
            content_type="multipart/form-data",
        )

    def test_upload_returns_file_id(self, client):
        resp = self._upload_txt(client)
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert "file_id" in data
        assert len(data["file_id"]) == 36  # UUID format

    def test_list_files_empty_initially(self, client):
        resp = client.get("/api/files")
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert "files" in data
        assert isinstance(data["files"], list)

    def test_upload_appears_in_file_list(self, client):
        upload = self._upload_txt(client, name="slides.txt")
        file_id = json.loads(upload.data)["file_id"]

        resp = client.get("/api/files")
        data = json.loads(resp.data)
        ids = [f["id"] for f in data["files"]]
        assert file_id in ids

    def test_get_file_by_id(self, client):
        upload = self._upload_txt(client, name="doc.txt")
        file_id = json.loads(upload.data)["file_id"]

        resp = client.get(f"/api/files/{file_id}")
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert data["file_id"] == file_id
        assert data["filename"] == "doc.txt"
        assert "slides" in data

    def test_get_nonexistent_file_404(self, client):
        resp = client.get("/api/files/00000000-0000-0000-0000-000000000000")
        assert resp.status_code == 404

    def test_delete_file(self, client):
        upload = self._upload_txt(client)
        file_id = json.loads(upload.data)["file_id"]

        del_resp = client.delete(f"/api/files/{file_id}")
        assert del_resp.status_code == 200

        # file should no longer appear in the list
        list_resp = client.get("/api/files")
        ids = [f["id"] for f in json.loads(list_resp.data)["files"]]
        assert file_id not in ids

    def test_delete_nonexistent_file_404(self, client):
        resp = client.delete("/api/files/00000000-0000-0000-0000-000000000000")
        assert resp.status_code == 404

    def test_file_metadata_fields(self, client):
        upload = self._upload_txt(client, name="meta.txt")
        file_id = json.loads(upload.data)["file_id"]

        resp = client.get("/api/files")
        data = json.loads(resp.data)
        entry = next(f for f in data["files"] if f["id"] == file_id)
        assert "filename" in entry
        assert "total_slides" in entry
        assert "created_at" in entry
        assert "thumbnail" in entry


# ─── ML learning API ─────────────────────────────────────────────────────


class TestMLLearning:
    """Tests for /api/learn and /api/suggestions endpoints."""

    def _learn(self, client, command="next_slide", text="next", lang="en", confidence=0.95):
        return client.post(
            "/api/learn",
            data=json.dumps({"command": command, "text": text, "lang": lang, "confidence": confidence}),
            content_type="application/json",
        )

    def test_learn_valid_command(self, client):
        resp = self._learn(client)
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert data.get("success") is True

    def test_learn_unknown_command_rejected(self, client):
        resp = self._learn(client, command="unknown")
        assert resp.status_code == 400

    def test_learn_empty_command_rejected(self, client):
        resp = client.post(
            "/api/learn",
            data=json.dumps({}),
            content_type="application/json",
        )
        assert resp.status_code == 400

    def test_suggestions_returns_list(self, client):
        # seed some commands
        self._learn(client, command="next_slide")
        self._learn(client, command="next_slide")
        self._learn(client, command="prev_slide")

        resp = client.get("/api/suggestions")
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert "suggestions" in data
        assert isinstance(data["suggestions"], list)

    def test_suggestions_sorted_by_frequency(self, client):
        # seed: next_slide ×3, prev_slide ×1
        for _ in range(3):
            self._learn(client, command="next_slide")
        self._learn(client, command="prev_slide")

        data = json.loads(client.get("/api/suggestions").data)
        top = data["suggestions"][0]["command"]
        assert top == "next_slide"

    def test_suggestions_limit_param(self, client):
        for cmd in ["next_slide", "prev_slide", "zoom_in", "zoom_out", "blackout", "fullscreen"]:
            self._learn(client, command=cmd)
        data = json.loads(client.get("/api/suggestions?limit=3").data)
        assert len(data["suggestions"]) <= 3

    def test_suggestion_has_required_fields(self, client):
        self._learn(client)
        data = json.loads(client.get("/api/suggestions").data)
        if data["suggestions"]:
            s = data["suggestions"][0]
            assert "command" in s
            assert "count" in s
            assert "avg_confidence" in s
            assert "last_used" in s


# ─── AI analysis helpers ──────────────────────────────────────────────────


class TestAIHelpers:
    """Unit tests for the AI text-analysis helper functions."""

    _TEXT = (
        "Machine learning is a powerful technology that improves performance. "
        "AI systems can analyze data and make intelligent decisions. "
        "This leads to significant improvements and great business growth."
    )

    def test_extract_keywords_returns_list(self):
        kws = _extract_keywords(self._TEXT)
        assert isinstance(kws, list)
        assert len(kws) > 0

    def test_extract_keywords_has_word_and_score(self):
        kws = _extract_keywords(self._TEXT, top_n=3)
        for kw in kws:
            assert "word" in kw
            assert "score" in kw
            assert 0.0 <= kw["score"] <= 1.0

    def test_extract_keywords_top_n_respected(self):
        kws = _extract_keywords(self._TEXT, top_n=4)
        assert len(kws) <= 4

    def test_extract_keywords_empty_text(self):
        kws = _extract_keywords("")
        assert kws == []

    def test_extractive_summary_shorter_than_original(self):
        long_text = (self._TEXT + " ") * 5  # repeat to create more sentences
        summary = _extractive_summary(long_text, num_sentences=2)
        assert len(summary) < len(long_text)

    def test_extractive_summary_non_empty(self):
        summary = _extractive_summary(self._TEXT)
        assert len(summary) > 0

    def test_extractive_summary_short_text_unchanged(self):
        short = "Hello world."
        summary = _extractive_summary(short, num_sentences=3)
        assert summary  # non-empty

    def test_sentiment_positive(self):
        assert _simple_sentiment("great success excellent improvement") == "positive"

    def test_sentiment_negative(self):
        assert _simple_sentiment("bad failure problem risk danger decline") == "negative"

    def test_sentiment_neutral(self):
        assert _simple_sentiment("the cat sat on the mat") == "neutral"

    def test_reading_time_positive(self):
        t = _estimate_reading_time(self._TEXT)
        assert t >= 1

    def test_reading_time_scales_with_length(self):
        short = "Hello world."
        long = (self._TEXT + " ") * 20
        assert _estimate_reading_time(long) > _estimate_reading_time(short)


# ─── AI analysis API route ────────────────────────────────────────────────


class TestAIRoute:
    """Tests for the /api/ai/analyze Flask endpoint."""

    def test_analyze_valid_text(self, client):
        resp = client.post(
            "/api/ai/analyze",
            data=json.dumps({"text": "Machine learning improves technology performance."}),
            content_type="application/json",
        )
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert "keywords" in data
        assert "summary" in data
        assert "sentiment" in data
        assert "reading_time_seconds" in data
        assert "word_count" in data

    def test_analyze_required_response_fields(self, client):
        resp = client.post(
            "/api/ai/analyze",
            data=json.dumps({"text": "Great success in AI and machine learning."}),
            content_type="application/json",
        )
        data = json.loads(resp.data)
        assert isinstance(data["keywords"], list)
        assert isinstance(data["summary"], str)
        assert data["sentiment"] in ("positive", "negative", "neutral")
        assert isinstance(data["reading_time_seconds"], int)
        assert isinstance(data["word_count"], int)

    def test_analyze_empty_text_returns_400(self, client):
        resp = client.post(
            "/api/ai/analyze",
            data=json.dumps({"text": ""}),
            content_type="application/json",
        )
        assert resp.status_code == 400
        data = json.loads(resp.data)
        assert "error" in data

    def test_analyze_missing_text_returns_400(self, client):
        resp = client.post(
            "/api/ai/analyze",
            data=json.dumps({}),
            content_type="application/json",
        )
        assert resp.status_code == 400

    def test_analyze_word_count_correct(self, client):
        text = "one two three four five"
        resp = client.post(
            "/api/ai/analyze",
            data=json.dumps({"text": text}),
            content_type="application/json",
        )
        data = json.loads(resp.data)
        assert data["word_count"] == 5

    def test_analyze_positive_sentiment(self, client):
        resp = client.post(
            "/api/ai/analyze",
            data=json.dumps({"text": "Excellent performance and great success achieved."}),
            content_type="application/json",
        )
        data = json.loads(resp.data)
        assert data["sentiment"] == "positive"

    def test_analyze_negative_sentiment(self, client):
        resp = client.post(
            "/api/ai/analyze",
            data=json.dumps({"text": "Bad failure and serious risk problem encountered."}),
            content_type="application/json",
        )
        data = json.loads(resp.data)
        assert data["sentiment"] == "negative"


# ─── root entrypoint ─────────────────────────────────────────────────────


class TestRootEntrypoint:
    """Smoke test for the root-level app.py entrypoint."""

    def test_root_entrypoint_imports_flask_app(self):
        import importlib.util
        from pathlib import Path

        spec = importlib.util.spec_from_file_location(
            "root_app",
            Path(__file__).resolve().parents[1] / "app.py",
        )
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        # The module must expose a Flask 'app' object
        assert hasattr(module, "app")
        from flask import Flask
        assert isinstance(module.app, Flask)




# ─── CSV upload ──────────────────────────────────────────────────────────────


class TestConvertCsv:
    """Unit tests for the convert_csv helper."""

    def _csv(self, text: str, name: str = "data.csv") -> bytes:
        return text.encode("utf-8")

    def test_basic_table(self):
        slides = convert_csv(self._csv("Name,Age\nAlice,30\nBob,25"), "data.csv")
        assert len(slides) == 1
        s = slides[0]
        assert s["type"] == "table"
        assert s["title"] == "data"
        assert s["headers"] == ["Name", "Age"]
        assert s["rows"] == [["Alice", "30"], ["Bob", "25"]]

    def test_empty_csv(self):
        slides = convert_csv(self._csv(""), "empty.csv")
        assert len(slides) == 1
        assert slides[0]["type"] == "table"
        assert slides[0]["headers"] == []

    def test_chunking(self):
        rows = ["h1,h2"] + [f"{i},{i*2}" for i in range(120)]
        slides = convert_csv(self._csv("\n".join(rows)), "big.csv")
        # 120 data rows → 3 slides of 50 / 50 / 20
        assert len(slides) == 3
        assert len(slides[0]["rows"]) == 50
        assert len(slides[1]["rows"]) == 50
        assert len(slides[2]["rows"]) == 20
        # Titles mention the chunk number
        assert "1/3" in slides[0]["title"]
        assert "3/3" in slides[2]["title"]

    def test_upload_csv_via_endpoint(self, client):
        csv_content = b"Product,Price\nApple,1.5\nBanana,0.9"
        data = {
            "file": (io.BytesIO(csv_content), "prices.csv"),
        }
        resp = client.post("/upload", data=data, content_type="multipart/form-data")
        assert resp.status_code == 200
        result = json.loads(resp.data)
        assert result["success"]
        assert result["filename"] == "prices.csv"
        assert result["total_slides"] == 1
        slide = result["slides"][0]
        assert slide["type"] == "table"
        assert slide["headers"] == ["Product", "Price"]


# ─── chart image analysis ─────────────────────────────────────────────────────


def _make_png_data_uri(width: int = 10, height: int = 10, color: tuple = (255, 0, 0)) -> str:
    """Create a minimal solid-color PNG as a data-URI for testing."""
    from PIL import Image
    img = Image.new("RGB", (width, height), color)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    b64 = __import__("base64").b64encode(buf.getvalue()).decode()
    return f"data:image/png;base64,{b64}"


class TestAnalyzeChartImage:
    """Unit tests for _analyze_chart_image and the /api/ai/analyze-chart endpoint."""

    def test_returns_expected_keys(self):
        uri = _make_png_data_uri()
        result = _analyze_chart_image(uri)
        for key in ("width", "height", "aspect_ratio", "brightness",
                    "colorfulness", "dominant_colors", "has_white_background",
                    "chart_type_hint"):
            assert key in result, f"Missing key: {key}"

    def test_dimensions(self):
        uri = _make_png_data_uri(width=80, height=40)
        result = _analyze_chart_image(uri)
        assert result["width"] == 80
        assert result["height"] == 40
        assert result["aspect_ratio"] == pytest.approx(2.0, rel=0.01)

    def test_bright_white_image(self):
        uri = _make_png_data_uri(width=20, height=20, color=(255, 255, 255))
        result = _analyze_chart_image(uri)
        assert result["brightness"] > 200
        assert result["has_white_background"] is True

    def test_dark_image(self):
        uri = _make_png_data_uri(width=20, height=20, color=(10, 10, 10))
        result = _analyze_chart_image(uri)
        assert result["brightness"] < 50
        assert result["has_white_background"] is False

    def test_dominant_colors_returned(self):
        uri = _make_png_data_uri(color=(100, 150, 200))
        result = _analyze_chart_image(uri)
        assert isinstance(result["dominant_colors"], list)
        assert len(result["dominant_colors"]) > 0
        for c in result["dominant_colors"]:
            assert c.startswith("#")
            assert len(c) == 7

    def test_chart_type_hint_is_string(self):
        uri = _make_png_data_uri()
        result = _analyze_chart_image(uri)
        valid_hints = {"bar", "pie", "line", "scatter", "unknown"}
        assert result["chart_type_hint"] in valid_hints

    def test_endpoint_no_image(self, client):
        resp = client.post(
            "/api/ai/analyze-chart",
            data=json.dumps({}),
            content_type="application/json",
        )
        assert resp.status_code == 400
        data = json.loads(resp.data)
        assert "error" in data

    def test_endpoint_with_image(self, client):
        uri = _make_png_data_uri(width=50, height=30, color=(200, 100, 50))
        resp = client.post(
            "/api/ai/analyze-chart",
            data=json.dumps({"image": uri}),
            content_type="application/json",
        )
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert data["width"] == 50
        assert data["height"] == 30
        assert "chart_type_hint" in data
        assert "dominant_colors" in data

    def test_endpoint_accepts_raw_base64(self, client):
        """Endpoint should also accept bare base64 (no data-URI prefix)."""
        from PIL import Image
        import base64
        img = Image.new("RGB", (10, 10), (0, 200, 0))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        raw_b64 = base64.b64encode(buf.getvalue()).decode()
        resp = client.post(
            "/api/ai/analyze-chart",
            data=json.dumps({"image": raw_b64}),
            content_type="application/json",
        )
        assert resp.status_code == 200
        data = json.loads(resp.data)
        assert data["width"] == 10
