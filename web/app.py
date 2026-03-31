#!/usr/bin/env python3
"""
Yot-Presentation Web Application
Converts uploaded files (PDF, Word, Excel, images, text) into
an online presentation with AI-powered voice command control.

New in v5.4:
- File management system: upload multiple files, switch between them, delete.
- ML learning: tracks voice command usage in SQLite; surfaces personalised
  suggestions to help frequent users work faster.
- Docker-ready: listens on 0.0.0.0, DATA_DIR env-var for the SQLite DB path.
"""

import base64
import io
import json
import os
import re
import sqlite3
import uuid
from datetime import datetime, timezone
from functools import wraps
from pathlib import Path
from typing import Any

from flask import Flask, jsonify, redirect, render_template, request, session, url_for
from flask_cors import CORS
from werkzeug.security import check_password_hash, generate_password_hash

app = Flask(__name__)
# Allow cross-origin requests from clients.  Set the
# ALLOWED_ORIGINS environment variable to a comma-separated list of
# origins to restrict access in production (e.g.
# "https://my-app.web.app,https://api.example.com").
# Defaults to "*" so local development works out of the box.
_raw_origins = os.environ.get("ALLOWED_ORIGINS", "*")
_cors_origins: list[str] | str = (
    [o.strip() for o in _raw_origins.split(",") if o.strip()]
    if _raw_origins != "*"
    else "*"
)
CORS(app, origins=_cors_origins)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB

# ─── runtime configuration ────────────────────────────────────────────────
DATA_DIR = Path(os.environ.get("DATA_DIR", Path(__file__).parent / "data"))
try:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
except PermissionError:
    # Serverless environments (e.g. Vercel) have a read-only deployment
    # filesystem; fall back to a writable temporary directory.
    import tempfile
    DATA_DIR = Path(tempfile.gettempdir()) / "yot-data"
    DATA_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = DATA_DIR / "yot_learning.db"

# ─── session secret key ───────────────────────────────────────────────────
# Use a stable key so sessions survive restarts.  Priority:
#  1. SECRET_KEY env-var (recommended for production)
#  2. .secret_key file persisted in DATA_DIR
_env_secret = os.environ.get("SECRET_KEY")
if _env_secret:
    app.secret_key = _env_secret
else:
    _secret_key_file = DATA_DIR / ".secret_key"
    if _secret_key_file.exists():
        app.secret_key = _secret_key_file.read_bytes()
    else:
        _generated_key = os.urandom(32)
        try:
            _secret_key_file.write_bytes(_generated_key)
        except OSError:
            pass
        app.secret_key = _generated_key

# In-memory file registry: file_id → {filename, total_slides, slides,
#                                      created_at, thumbnail}
# NOTE: the file registry is EPHEMERAL — it resets on server/container restart.
# The ML learning database (yot_learning.db) is PERSISTENT because it is
# written to DATA_DIR, which is mounted as a Docker named volume.
_file_registry: dict[str, dict[str, Any]] = {}

ALLOWED_EXTENSIONS = {
    "pdf",
    "docx",
    "doc",
    "xlsx",
    "xls",
    "csv",
    "txt",
    "png",
    "jpg",
    "jpeg",
    "gif",
    "bmp",
    "webp",
}

# ─────────────────────────────────────────────
# helpers
# ─────────────────────────────────────────────


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


# ─────────────────────────────────────────────
# file → slides converters
# ─────────────────────────────────────────────


def convert_pdf(file_bytes: bytes) -> list[dict[str, Any]]:
    """Render each PDF page as a PNG image slide."""
    import fitz  # PyMuPDF

    slides: list[dict[str, Any]] = []
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    for page_num, page in enumerate(doc):
        mat = fitz.Matrix(1.5, 1.5)
        pix = page.get_pixmap(matrix=mat)
        img_b64 = base64.b64encode(pix.tobytes("png")).decode()
        slides.append(
            {
                "id": page_num + 1,
                "type": "image",
                "title": f"Page {page_num + 1}",
                "image": f"data:image/png;base64,{img_b64}",
                "notes": page.get_text().strip(),
            }
        )
    doc.close()
    return slides


def convert_docx(file_bytes: bytes) -> list[dict[str, Any]]:
    """Convert a Word document to text/bullet slides grouped by Heading 1."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    slide_groups: list[dict[str, Any]] = []
    current: dict[str, Any] = {"title": "", "bullets": []}
    found_heading1 = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower()
        if "heading 1" in style:
            if current["title"] or current["bullets"]:
                slide_groups.append(current)
            current = {"title": text, "bullets": []}
            found_heading1 = True
        elif "heading 2" in style or "heading 3" in style:
            current["bullets"].append({"level": 1, "text": text})
        else:
            current["bullets"].append({"level": 2, "text": text})

    if current["title"] or current["bullets"]:
        slide_groups.append(current)

    # Fallback – no Heading 1 found; chunk all paragraphs into slides of ≤ 6 lines
    if not found_heading1:
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        chunk = 6
        slide_groups = []
        for i in range(0, len(lines), chunk):
            block = lines[i : i + chunk]
            slide_groups.append(
                {
                    "title": block[0],
                    "bullets": [{"level": 2, "text": l} for l in block[1:]],
                }
            )

    return [
        {
            "id": i + 1,
            "type": "text",
            "title": g.get("title", f"Slide {i + 1}"),
            "bullets": g.get("bullets", []),
        }
        for i, g in enumerate(slide_groups)
    ]


def convert_xlsx(file_bytes: bytes) -> list[dict[str, Any]]:
    """Convert each Excel sheet to a table slide."""
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    slides: list[dict[str, Any]] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = [
            [str(cell) if cell is not None else "" for cell in row]
            for row in ws.iter_rows(values_only=True)
            if any(c is not None for c in row)
        ]
        if rows:
            slides.append(
                {
                    "id": len(slides) + 1,
                    "type": "table",
                    "title": sheet_name,
                    "headers": rows[0],
                    "rows": rows[1:],
                }
            )

    return slides


def convert_text(file_bytes: bytes) -> list[dict[str, Any]]:
    """Split a plain-text document on blank lines / form-feeds into slides."""
    text = file_bytes.decode("utf-8", errors="replace")
    sections = [s.strip() for s in re.split(r"\n\s*\n|\f", text) if s.strip()]

    if not sections:
        return [{"id": 1, "type": "text", "title": "Document", "bullets": []}]

    slides = []
    for i, section in enumerate(sections):
        lines = section.split("\n")
        slides.append(
            {
                "id": i + 1,
                "type": "text",
                "title": lines[0],
                "bullets": [
                    {"level": 2, "text": l.strip()}
                    for l in lines[1:]
                    if l.strip()
                ],
            }
        )
    return slides


def convert_image(file_bytes: bytes, filename: str) -> list[dict[str, Any]]:
    """Wrap a single image as one slide."""
    ext = filename.rsplit(".", 1)[-1].lower()
    mime = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"
    img_b64 = base64.b64encode(file_bytes).decode()
    return [
        {
            "id": 1,
            "type": "image",
            "title": Path(filename).stem,
            "image": f"data:{mime};base64,{img_b64}",
        }
    ]


def convert_csv(file_bytes: bytes, filename: str) -> list[dict[str, Any]]:
    """Convert a CSV file into table slides (one slide per 50 rows)."""
    import csv

    text = file_bytes.decode("utf-8-sig", errors="replace")
    reader = csv.reader(text.splitlines())
    rows_raw = [row for row in reader if any(cell.strip() for cell in row)]
    if not rows_raw:
        return [{"id": 1, "type": "table", "title": Path(filename).stem, "headers": [], "rows": []}]

    headers = rows_raw[0]
    data_rows = rows_raw[1:]

    # Chunk into slides of at most 50 data rows each
    chunk_size = 50
    slides: list[dict[str, Any]] = []
    chunks = [data_rows[i : i + chunk_size] for i in range(0, len(data_rows), chunk_size)] if data_rows else [[]]
    stem = Path(filename).stem
    for i, chunk in enumerate(chunks):
        title = stem if len(chunks) == 1 else f"{stem} ({i + 1}/{len(chunks)})"
        slides.append(
            {
                "id": i + 1,
                "type": "table",
                "title": title,
                "headers": headers,
                "rows": chunk,
            }
        )
    return slides


# ─────────────────────────────────────────────
# voice-command matching (mirrors original Python logic)
# ─────────────────────────────────────────────

# Multi-language command patterns identical to the original v5.3.1 application
COMMAND_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    # jump to slide (must come first – has a capture group)
    (
        "jump_slide",
        re.compile(
            r"(?:jump to|go to|slide|page|number|"
            r"salta a|ve a|diapositiva|página|"
            r"aller à|diapo|numéro|"
            r"gehe zu|folie|seite|"
            r"vai a|ir para|"
            r"跳到|转到|幻灯片|スライド|ページ)"
            r"\s*(\d+)",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # next slide
    (
        "next_slide",
        re.compile(
            r"\b(next|forward|advance|go forward|go right|"
            r"siguiente|adelante|próxima|"
            r"suivant|avancer|prochaine|"
            r"nächst|vorwärts|"
            r"prossimo|avanti|successivo|"
            r"próximo|avançar|seguinte|"
            r"下一张|下一个|向前|次へ|進む)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # previous slide
    (
        "prev_slide",
        re.compile(
            r"\b(previous|prev|back|go back|return|"
            r"anterior|atrás|volver|"
            r"précédent|retour|"
            r"zurück|vorherig|"
            r"precedente|indietro|tornare|"
            r"voltar|para trás|"
            r"上一张|上一个|向后|戻る|前へ)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # start presentation
    (
        "start_show",
        re.compile(
            r"\b(start presentation|begin show|present now|"
            r"comenzar presentación|iniciar show|"
            r"commencer présentation|débuter|"
            r"präsentation starten|"
            r"inizia presentazione|"
            r"iniciar apresentação|"
            r"开始演示|开始放映|"
            r"プレゼンテーション開始|スライドショー開始)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # start (shorter fallback)
    (
        "start_show",
        re.compile(
            r"\b(start|begin|present|play|iniciar|commencer|starten|iniziare|começar)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # end show
    (
        "end_show",
        re.compile(
            r"\b(stop presentation|end show|exit show|close presentation|"
            r"detener presentación|finalizar show|"
            r"arrêter présentation|quitter diaporama|"
            r"präsentation beenden|"
            r"ferma presentazione|termina spettacolo|"
            r"parar apresentação|sair do show|"
            r"停止演示|退出幻灯片|"
            r"プレゼンテーション終了|スライドショー終了)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # end (shorter fallback)
    (
        "end_show",
        re.compile(
            r"\b(end|stop|exit|quit|close|terminar|finir|beenden|terminare)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # blackout
    (
        "blackout",
        re.compile(
            r"\b(black|blackout|blank|darken|turn off screen|"
            r"pantalla negra|oscurecer|"
            r"écran noir|assombrir|"
            r"schwarzer bildschirm|verdunkeln|"
            r"schermo nero|scurire|"
            r"tela preta|escurecer|"
            r"黑屏|黒い画面|暗くする)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # zoom in
    (
        "zoom_in",
        re.compile(
            r"\b(zoom in|magnify|enlarge|agrandir|vergrößern|ingrandire|ampliar|放大|ズームイン)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # zoom out
    (
        "zoom_out",
        re.compile(
            r"\b(zoom out|shrink|reduce|réduire|verkleinern|rimpicciolire|縮小|ズームアウト)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # zoom reset
    (
        "zoom_reset",
        re.compile(
            r"\b(reset zoom|normal size|actual size|zoom normal|rétablir zoom)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # fullscreen
    (
        "fullscreen",
        re.compile(
            r"\b(fullscreen|full screen|maximize|présentation plein écran)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # first slide
    (
        "first_slide",
        re.compile(
            r"\b(first slide|go to start|beginning|primera|première|erste folie)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # last slide
    (
        "last_slide",
        re.compile(
            r"\b(last slide|final slide|end slide|última|dernière|letzte folie)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # pen tool
    (
        "pen_tool",
        re.compile(
            r"\b(pen tool|draw|annotation|herramienta pluma|dibujar|"
            r"outil stylo|dessiner|stiftwerkzeug|zeichnen|"
            r"strumento penna|disegnare|ferramenta caneta|desenhar|"
            r"笔工具|绘制|ペンツール|描画)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # eraser
    (
        "eraser",
        re.compile(
            r"\b(eraser|erase|clear drawing|borrar|gomme|effacer|radiergummi|gomma)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # pointer
    (
        "pointer",
        re.compile(
            r"\b(pointer|arrow|cursor|puntero|pointeur|zeiger|puntatore)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
]


def match_command(text: str) -> dict[str, Any]:
    """
    Match a spoken phrase against all command patterns.
    Returns a dict with at minimum an 'action' key.
    Mirrors the two-stage (regex → fuzzy) logic of the original v5.3.1.
    """
    text = text.lower().strip()

    for action, pattern in COMMAND_PATTERNS:
        m = pattern.search(text)
        if m:
            result: dict[str, Any] = {"action": action, "confidence": 0.95}
            if action == "jump_slide":
                result["slide"] = int(m.group(1))
            return result

    return {"action": "unknown", "text": text, "confidence": 0.0}


# ─────────────────────────────────────────────
# ML learning – SQLite-backed command frequency
# ─────────────────────────────────────────────


def _get_db() -> sqlite3.Connection:
    """Return a connection to the learning database, creating tables if needed."""
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS command_usage (
            id        INTEGER PRIMARY KEY AUTOINCREMENT,
            command   TEXT    NOT NULL,
            text      TEXT    NOT NULL,
            lang      TEXT    NOT NULL DEFAULT 'en',
            confidence REAL   NOT NULL DEFAULT 0.0,
            ts        TEXT    NOT NULL
        )
        """
    )
    conn.commit()
    return conn


def record_command(command: str, text: str, lang: str, confidence: float) -> None:
    """Insert one command-usage record into the learning database."""
    ts = datetime.now(timezone.utc).isoformat()
    with _get_db() as conn:
        conn.execute(
            "INSERT INTO command_usage (command, text, lang, confidence, ts) VALUES (?,?,?,?,?)",
            (command, text, lang, confidence, ts),
        )


def get_suggestions(limit: int = 5) -> list[dict[str, Any]]:
    """Return the most-frequently-used commands, highest count first."""
    with _get_db() as conn:
        rows = conn.execute(
            """
            SELECT command,
                   COUNT(*)                             AS count,
                   AVG(confidence)                      AS avg_confidence,
                   MAX(ts)                              AS last_used
            FROM   command_usage
            WHERE  command != 'unknown'
            GROUP  BY command
            ORDER  BY count DESC, last_used DESC
            LIMIT  ?
            """,
            (limit,),
        ).fetchall()
    return [
        {
            "command": row["command"],
            "count": row["count"],
            "avg_confidence": round(row["avg_confidence"], 3),
            "last_used": row["last_used"],
        }
        for row in rows
    ]


# ─────────────────────────────────────────────
# admin – authentication & event tracking
# ─────────────────────────────────────────────


def _get_admin_db() -> sqlite3.Connection:
    """Return a DB connection with admin tables created if not present."""
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    conn.executescript(
        """
        CREATE TABLE IF NOT EXISTS admin_users (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            username     TEXT    UNIQUE NOT NULL,
            password_hash TEXT   NOT NULL,
            created_at   TEXT   NOT NULL
        );
        CREATE TABLE IF NOT EXISTS site_events (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            event_type TEXT    NOT NULL,
            detail     TEXT    NOT NULL DEFAULT '',
            ts         TEXT    NOT NULL
        );
        """
    )
    conn.commit()
    return conn


def _admin_exists() -> bool:
    """Return True when at least one admin account has been registered."""
    with _get_admin_db() as conn:
        row = conn.execute("SELECT COUNT(*) AS cnt FROM admin_users").fetchone()
        return row["cnt"] > 0


def _record_event(event_type: str, detail: str = "") -> None:
    """Persist a site event for the admin dashboard."""
    ts = datetime.now(timezone.utc).isoformat()
    with _get_admin_db() as conn:
        conn.execute(
            "INSERT INTO site_events (event_type, detail, ts) VALUES (?,?,?)",
            (event_type, detail, ts),
        )


def _admin_login_required(f):
    """Decorator that redirects unauthenticated requests to the login page."""
    @wraps(f)
    def _decorated(*args, **kwargs):
        if not session.get("admin_logged_in"):
            return redirect(url_for("admin_login"))
        return f(*args, **kwargs)
    return _decorated


# ─────────────────────────────────────────────
# file management helpers
# ─────────────────────────────────────────────


def _thumbnail_for_slides(slides: list[dict[str, Any]]) -> str:
    """Return a data-URI thumbnail from the first slide (image type only)."""
    if slides and slides[0].get("type") == "image":
        return slides[0]["image"]
    return ""


def _register_file(
    filename: str, slides: list[dict[str, Any]]
) -> str:
    """Store the file in the registry and return its UUID."""
    file_id = str(uuid.uuid4())
    _file_registry[file_id] = {
        "id": file_id,
        "filename": filename,
        "total_slides": len(slides),
        "slides": slides,
        "thumbnail": _thumbnail_for_slides(slides),
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    return file_id


# ─────────────────────────────────────────────
# routes
# ─────────────────────────────────────────────


@app.route("/")
def index():
    try:
        return render_template("index.html")
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "No file selected"}), 400

    if not allowed_file(file.filename):
        return (
            jsonify(
                {
                    "error": (
                        f"Unsupported file type. "
                        f"Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
                    )
                }
            ),
            400,
        )

    file_bytes = file.read()
    filename: str = file.filename
    ext = filename.rsplit(".", 1)[1].lower()

    try:
        converters = {
            "pdf": lambda b: convert_pdf(b),
            "docx": lambda b: convert_docx(b),
            "doc": lambda b: convert_docx(b),
            "xlsx": lambda b: convert_xlsx(b),
            "xls": lambda b: convert_xlsx(b),
            "csv": lambda b: convert_csv(b, filename),
            "txt": lambda b: convert_text(b),
            "png": lambda b: convert_image(b, filename),
            "jpg": lambda b: convert_image(b, filename),
            "jpeg": lambda b: convert_image(b, filename),
            "gif": lambda b: convert_image(b, filename),
            "bmp": lambda b: convert_image(b, filename),
            "webp": lambda b: convert_image(b, filename),
        }
        slides = converters[ext](file_bytes)
        file_id = _register_file(filename, slides)
        _record_event("upload", filename)
        return jsonify(
            {
                "success": True,
                "file_id": file_id,
                "filename": filename,
                "total_slides": len(slides),
                "slides": slides,
            }
        )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ─── file management ─────────────────────────────────────────────────────────


@app.route("/api/files", methods=["GET"])
def list_files():
    """Return metadata for all uploaded files (no slide data to keep payload small)."""
    try:
        files = [
            {
                "id": entry["id"],
                "filename": entry["filename"],
                "total_slides": entry["total_slides"],
                "thumbnail": entry["thumbnail"],
                "created_at": entry["created_at"],
            }
            for entry in _file_registry.values()
        ]
        # newest first
        files.sort(key=lambda f: f["created_at"], reverse=True)
        return jsonify({"files": files})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/files/<file_id>", methods=["GET"])
def get_file(file_id: str):
    """Return the full slide data for a single file."""
    try:
        entry = _file_registry.get(file_id)
        if entry is None:
            return jsonify({"error": "File not found"}), 404
        return jsonify(
            {
                "success": True,
                "file_id": file_id,
                "filename": entry["filename"],
                "total_slides": entry["total_slides"],
                "slides": entry["slides"],
            }
        )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/files/<file_id>", methods=["DELETE"])
def delete_file(file_id: str):
    """Remove a file from the in-memory registry."""
    try:
        if file_id not in _file_registry:
            return jsonify({"error": "File not found"}), 404
        filename = _file_registry[file_id]["filename"]
        del _file_registry[file_id]
        _record_event("delete", filename)
        return jsonify({"success": True, "deleted": file_id})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ─── ML learning ─────────────────────────────────────────────────────────────


@app.route("/api/learn", methods=["POST"])
def learn():
    """Record a successfully executed voice command for ML training."""
    data: dict[str, Any] = request.get_json(force=True) or {}
    command: str = data.get("command", "")
    text: str = data.get("text", "")
    lang: str = data.get("lang", "en")
    confidence: float = float(data.get("confidence", 0.0))

    if not command or command == "unknown":
        return jsonify({"error": "No valid command provided"}), 400

    try:
        record_command(command, text, lang, confidence)
        _record_event("voice_command", command)
        return jsonify({"success": True})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/suggestions", methods=["GET"])
def suggestions():
    """Return the user's most-frequently-used commands (ML-derived)."""
    limit = min(int(request.args.get("limit", 5)), 20)
    try:
        return jsonify({"suggestions": get_suggestions(limit)})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/command", methods=["POST"])
def process_command():
    """Process a voice-command transcript and return the matched action."""
    try:
        data: dict[str, Any] = request.get_json(force=True) or {}
        text: str = data.get("text", "")
        result = match_command(text)
        return jsonify(result)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ─── AI text analysis ─────────────────────────────────────────────────────


def _tokenize(text: str) -> list[str]:
    """Return a list of lowercase alphabetic tokens, filtering stopwords."""
    _STOPWORDS = {
        "a", "an", "the", "and", "or", "but", "in", "on", "at", "to",
        "for", "of", "with", "by", "from", "is", "are", "was", "were",
        "be", "been", "being", "have", "has", "had", "do", "does", "did",
        "will", "would", "shall", "should", "may", "might", "must", "can",
        "could", "this", "that", "these", "those", "it", "its", "as", "if",
        "not", "no", "so", "than", "then", "when", "where", "which", "who",
        "what", "how", "all", "each", "more", "also", "about", "up", "out",
        "into", "over", "after", "before", "i", "we", "you", "he", "she",
        "they", "their", "our", "your", "his", "her", "my",
    }
    tokens = re.findall(r"[a-z]{3,}", text.lower())
    return [t for t in tokens if t not in _STOPWORDS]


def _extract_keywords(text: str, top_n: int = 8) -> list[dict[str, Any]]:
    """Return the top-N most frequent content words with normalised scores."""
    tokens = _tokenize(text)
    if not tokens:
        return []
    freq: dict[str, int] = {}
    for tok in tokens:
        freq[tok] = freq.get(tok, 0) + 1
    max_freq = max(freq.values())
    ranked = sorted(freq.items(), key=lambda x: x[1], reverse=True)[:top_n]
    return [{"word": w, "score": round(c / max_freq, 3)} for w, c in ranked]


def _extractive_summary(text: str, num_sentences: int = 3) -> str:
    """Simple frequency-based extractive summarization."""
    sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if len(s.strip()) > 20]
    if not sentences:
        return text[:300] if text else ""
    if len(sentences) <= num_sentences:
        return " ".join(sentences)

    tokens = _tokenize(text)
    freq: dict[str, int] = {}
    for tok in tokens:
        freq[tok] = freq.get(tok, 0) + 1

    def score_sentence(s: str) -> float:
        words = _tokenize(s)
        return sum(freq.get(w, 0) for w in words) / max(len(words), 1)

    ranked = sorted(sentences, key=score_sentence, reverse=True)[:num_sentences]
    # Preserve original order
    ordered = [s for s in sentences if s in ranked]
    return " ".join(ordered)


def _estimate_reading_time(text: str) -> int:
    """Estimate reading time in seconds at 200 words per minute."""
    word_count = len(text.split())
    return max(1, round(word_count / 200 * 60))


def _simple_sentiment(text: str) -> str:
    """Classify text sentiment as positive, negative, or neutral."""
    _POS = {"good", "great", "excellent", "success", "improve", "benefit",
             "positive", "increase", "growth", "effective", "best", "advantage",
             "opportunity", "achieve", "win", "progress", "innovation", "strong"}
    _NEG = {"bad", "fail", "problem", "issue", "risk", "negative", "decrease",
             "loss", "poor", "weak", "challenge", "threat", "error", "wrong",
             "difficult", "decline", "concern", "crisis", "danger"}
    words = set(_tokenize(text))
    pos = len(words & _POS)
    neg = len(words & _NEG)
    if pos > neg:
        return "positive"
    if neg > pos:
        return "negative"
    return "neutral"


@app.route("/api/ai/analyze", methods=["POST"])
def ai_analyze():
    """
    Analyze the text content of one or more slides and return:
     - keywords  : top frequent content words with relative scores
     - summary   : extractive summary (3 sentences max)
     - sentiment : positive / negative / neutral
     - reading_time_seconds : estimated reading time
     - word_count : total word count
    """
    data: dict[str, Any] = request.get_json(force=True) or {}
    text: str = data.get("text", "").strip()

    if not text:
        return jsonify({"error": "No text provided"}), 400

    try:
        return jsonify({
            "keywords": _extract_keywords(text),
            "summary": _extractive_summary(text),
            "sentiment": _simple_sentiment(text),
            "reading_time_seconds": _estimate_reading_time(text),
            "word_count": len(text.split()),
        })
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ─── AI chart image analysis ──────────────────────────────────────────────


def _analyze_chart_image(image_data_uri: str) -> dict[str, Any]:
    """
    Analyze a chart image supplied as a data-URI and return visual statistics.

    Uses Pillow to compute:
     - width / height / aspect_ratio
     - brightness  : average perceived luminance (0–255)
     - colorfulness : standard deviation of hue across pixels (0–180)
     - dominant_colors : up to 6 hex colors obtained via median-cut quantization
     - has_white_background : True when the majority of edge pixels are near-white
     - chart_type_hint : a best-effort guess (bar / pie / line / scatter / unknown)
       based on simple heuristic analysis of dominant colors and aspect ratio
    """
    from PIL import Image, ImageStat
    import colorsys

    # Strip data-URI prefix and decode
    if "," in image_data_uri:
        image_data_uri = image_data_uri.split(",", 1)[1]
    img_bytes = base64.b64decode(image_data_uri)
    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")

    width, height = img.size
    aspect_ratio = round(width / height, 3) if height else 1.0

    # Brightness: mean of luminance channel from L*a*b* would be ideal, but
    # converting via YCbCr gives a good perceptual estimate with Pillow alone.
    lum_img = img.convert("L")
    stat = ImageStat.Stat(lum_img)
    brightness = round(stat.mean[0], 1)

    # Colorfulness: measure hue spread across a small grid of sampled pixels
    sample_img = img.resize((64, 64), Image.LANCZOS)
    # Use get_flattened_data when available (Pillow ≥ 10 deprecates getdata)
    _getter = getattr(sample_img, "get_flattened_data", None) or sample_img.getdata
    pixels = list(_getter())
    hues: list[float] = []
    for r, g, b in pixels:
        h, s, v = colorsys.rgb_to_hsv(r / 255, g / 255, b / 255)
        if s > 0.15:  # only count pixels that are actually coloured
            hues.append(h * 360)
    colorfulness = 0.0
    if hues:
        mean_h = sum(hues) / len(hues)
        colorfulness = round((sum((h - mean_h) ** 2 for h in hues) / len(hues)) ** 0.5, 1)

    # Dominant colors via Pillow's built-in quantization
    palette_img = img.convert("P", palette=Image.ADAPTIVE, colors=6)
    palette = palette_img.getpalette()  # flat list R,G,B,R,G,B,…
    dominant_colors: list[str] = []
    if palette:
        num_colors = min(6, len(palette) // 3)
        for i in range(num_colors):
            r, g, b = palette[i * 3], palette[i * 3 + 1], palette[i * 3 + 2]
            dominant_colors.append(f"#{r:02x}{g:02x}{b:02x}")

    # White-background heuristic: sample a 1-pixel border; if ≥ 70% of those
    # pixels have all channels ≥ 230 we call it a white background.
    border_pixels: list[tuple[int, int, int]] = []
    arr = img.load()
    for x in range(width):
        border_pixels.append(arr[x, 0])
        border_pixels.append(arr[x, height - 1])
    for y in range(height):
        border_pixels.append(arr[0, y])
        border_pixels.append(arr[width - 1, y])
    white_count = sum(1 for r, g, b in border_pixels if r >= 230 and g >= 230 and b >= 230)
    has_white_background = bool(border_pixels and white_count / len(border_pixels) >= 0.7)

    # Chart-type heuristic
    chart_type_hint = "unknown"
    if colorfulness < 20 and len(hues) < 50:
        chart_type_hint = "line"
    elif colorfulness >= 20 and aspect_ratio > 1.2:
        chart_type_hint = "bar"
    elif 0.8 <= aspect_ratio <= 1.2 and colorfulness >= 30:
        chart_type_hint = "pie"
    elif colorfulness >= 15 and brightness > 180:
        chart_type_hint = "scatter"
    elif colorfulness >= 20:
        chart_type_hint = "bar"

    return {
        "width": width,
        "height": height,
        "aspect_ratio": aspect_ratio,
        "brightness": brightness,
        "colorfulness": colorfulness,
        "dominant_colors": dominant_colors,
        "has_white_background": has_white_background,
        "chart_type_hint": chart_type_hint,
    }


@app.route("/api/ai/analyze-chart", methods=["POST"])
def ai_analyze_chart():
    """
    Analyze a chart image and return visual statistics.

    Request body (JSON):
      { "image": "<data-URI of chart image>" }

    Response:
      {
        "width": int, "height": int, "aspect_ratio": float,
        "brightness": float,
        "colorfulness": float,
        "dominant_colors": ["#rrggbb", …],
        "has_white_background": bool,
        "chart_type_hint": "bar" | "pie" | "line" | "scatter" | "unknown"
      }
    """
    data: dict[str, Any] = request.get_json(force=True) or {}
    image: str = data.get("image", "").strip()

    if not image:
        return jsonify({"error": "No image data provided"}), 400

    try:
        result = _analyze_chart_image(image)
        return jsonify(result)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ─── admin authentication ─────────────────────────────────────────────────────


@app.route("/admin")
def admin_index():
    """Redirect to dashboard if logged in, otherwise to login."""
    if session.get("admin_logged_in"):
        return redirect(url_for("admin_dashboard"))
    return redirect(url_for("admin_login"))


@app.route("/admin/login", methods=["GET", "POST"])
def admin_login():
    """Login page for the admin.  Shows a registration form when no admin exists yet."""
    error = None
    if request.method == "POST":
        action = request.form.get("action", "login")
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")

        if action == "register":
            if _admin_exists():
                error = "An admin account already exists. Please log in."
            elif not username or not password:
                error = "Username and password are required."
            elif len(password) < 8:
                error = "Password must be at least 8 characters."
            else:
                ts = datetime.now(timezone.utc).isoformat()
                with _get_admin_db() as conn:
                    conn.execute(
                        "INSERT INTO admin_users (username, password_hash, created_at) VALUES (?,?,?)",
                        (username, generate_password_hash(password), ts),
                    )
                session["admin_logged_in"] = True
                session["admin_username"] = username
                _record_event("admin_register", username)
                return redirect(url_for("admin_dashboard"))
        else:
            # login
            with _get_admin_db() as conn:
                row = conn.execute(
                    "SELECT password_hash FROM admin_users WHERE username = ?",
                    (username,),
                ).fetchone()
            if row and check_password_hash(row["password_hash"], password):
                session["admin_logged_in"] = True
                session["admin_username"] = username
                _record_event("admin_login", username)
                return redirect(url_for("admin_dashboard"))
            error = "Invalid username or password."

    registration_open = not _admin_exists()
    return render_template(
        "admin_login.html",
        error=error,
        registration_open=registration_open,
    )


@app.route("/admin/logout")
def admin_logout():
    """Clear the admin session and redirect to login."""
    session.clear()
    return redirect(url_for("admin_login"))


@app.route("/admin/dashboard")
@_admin_login_required
def admin_dashboard():
    """Admin dashboard – overview page."""
    return render_template(
        "admin_dashboard.html",
        username=session.get("admin_username", "Admin"),
    )


@app.route("/admin/api/stats")
@_admin_login_required
def admin_stats():
    """Return JSON statistics for the admin dashboard."""
    try:
        with _get_admin_db() as conn:
            # Event counts by type
            event_rows = conn.execute(
                """
                SELECT event_type, COUNT(*) AS cnt
                FROM   site_events
                GROUP  BY event_type
                ORDER  BY cnt DESC
                """
            ).fetchall()
            event_counts: dict[str, int] = {r["event_type"]: r["cnt"] for r in event_rows}

            # Recent events (last 50)
            recent_rows = conn.execute(
                """
                SELECT event_type, detail, ts
                FROM   site_events
                ORDER  BY id DESC
                LIMIT  50
                """
            ).fetchall()
            recent_events = [
                {"event_type": r["event_type"], "detail": r["detail"], "ts": r["ts"]}
                for r in recent_rows
            ]

        # Voice-command frequency from ML learning DB
        top_commands = get_suggestions(limit=10)

        # Daily uploads – last 14 days
        with _get_admin_db() as conn:
            daily_rows = conn.execute(
                """
                SELECT substr(ts, 1, 10) AS day, COUNT(*) AS cnt
                FROM   site_events
                WHERE  event_type = 'upload'
                  AND  ts >= datetime('now', '-14 days')
                GROUP  BY day
                ORDER  BY day
                """
            ).fetchall()
            daily_uploads = [{"day": r["day"], "count": r["cnt"]} for r in daily_rows]

        return jsonify(
            {
                "totals": {
                    "uploads": event_counts.get("upload", 0),
                    "deletes": event_counts.get("delete", 0),
                    "voice_commands": event_counts.get("voice_command", 0),
                    "files_in_library": len(_file_registry),
                },
                "top_commands": top_commands,
                "recent_events": recent_events,
                "daily_uploads": daily_uploads,
            }
        )
    except Exception:
        return jsonify({"error": "Failed to load statistics"}), 500
    port = int(os.environ.get("PORT", 5000))
    debug = os.environ.get("FLASK_DEBUG", "0") == "1"
    app.run(debug=debug, host="0.0.0.0", port=port)
